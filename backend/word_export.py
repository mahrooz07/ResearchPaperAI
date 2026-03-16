from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_document(title: str, authors: list, sections: list, references: list, images: list, output_filename: str = "paper.docx") -> str:
    document = Document()

    # Title
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)

    # Authors
    for author in authors:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        name_run = p.add_run(author.name + "\n")
        name_run.bold = True
        
        affiliation_run = p.add_run(f"{author.department}\n{author.affiliation}\n{author.country}\n{author.email}")
        affiliation_run.font.size = Pt(10)

    document.add_page_break()

    # Sections (including abstract etc)
    for section in sections:
        heading = document.add_heading(section['section'], level=1)
        
        # Split content by logical paragraphs (simple newline split)
        content_paragraphs = section['content'].split('\\n')
        for para_text in content_paragraphs:
            if para_text.strip():
                p = document.add_paragraph(para_text.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
        # Insert image immediately under the section if requested
        if section.get('image_provided'):
            if section.get('image_path'):
                filename = os.path.basename(section['image_path'])
                local_path = os.path.join("uploads", filename)
                if os.path.exists(local_path):
                    document.add_picture(local_path, width=Inches(6.0))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    document.add_paragraph(f"[Image not found: {filename}]")
            else:
                p = document.add_paragraph(f"(Insert diagram for {section['section']} here)")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style.font.italic = True

    # Global Images Catch-all
    unused_images = [img for img in images if not any(img == s.get('image_path') for s in sections)]
    if unused_images:
        document.add_heading("Appendix: Figures", level=1)
        for img_url in unused_images:
            filename = os.path.basename(img_url)
            local_path = os.path.join("uploads", filename)
            if os.path.exists(local_path):
                document.add_picture(local_path, width=Inches(6.0))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                document.add_paragraph(f"[Image not found: {filename}]")

    # References
    if references:
        document.add_heading("References", level=1)
        for idx, ref in enumerate(references):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-20)
            p.add_run(f"[{idx+1}] {ref}")

    # Ensure uploads dir exists for saving
    os.makedirs("uploads", exist_ok=True)
    output_path = os.path.join("uploads", output_filename)
    document.save(output_path)
    
    return f"/uploads/{output_filename}"
